# utils/document_generator.py

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os
from datetime import datetime
from database.db import get_notification_with_details, get_session
from database.models import NotificationMeta, NotificationConsultation, Subject

def generate_document(notification_id):
    """Генерирует документ уведомления на основе данных из БД и шаблона"""
    # Получаем данные уведомления
    notification_data = get_notification_with_details(notification_id)
    
    if not notification_data:
        return {'success': False, 'message': 'Уведомление не найдено'}
    
    # Путь к единому шаблону
    template_path = 'templates/notification.docx'
    
    if not os.path.exists(template_path):
        return {'success': False, 'message': f'Шаблон не найден: {template_path}'}
    
    try:
        # Загружаем шаблон
        doc = Document(template_path)
        
        # Получаем данные для подстановки
        student = notification_data['student']
        subjects = notification_data['subjects']
        
        # Получаем дополнительные метаданные и консультации с предварительной загрузкой связанных данных
        session = get_session()
        meta_records = session.query(NotificationMeta).filter_by(notification_id=notification_id).all()
        meta = {item.key: item.value for item in meta_records}
        
        # Предзагружаем связанные объекты предметов в рамках одной сессии
        from sqlalchemy.orm import joinedload
        consultations_query = session.query(NotificationConsultation).options(
            joinedload(NotificationConsultation.subject)
        ).filter_by(notification_id=notification_id).all()
        
        # Создаем словарь с полной информацией о консультациях, включая имя предмета
        consultations = []
        for consultation in consultations_query:
            consultations.append({
                'subject_name': consultation.subject.name,
                'topic_name': consultation.topic_name,
                'date': consultation.date,
                'time': consultation.time,
                'topic_type': consultation.topic_type
            })
        
        # Теперь можно закрыть сессию, так как все данные уже получены
        session.close()
        
        # Определяем типы предметов
        failed_subjects = []
        satisfactory_subjects = []
        
        if 'failed_subjects' in meta:
            failed_subject_names = meta['failed_subjects'].split(',')
            failed_subjects = [s for s in subjects if s.name in failed_subject_names]
        else:
            # Если нет метаданных, все предметы считаем с задолженностями
            failed_subjects = subjects
        
        if 'satisfactory_subjects' in meta:
            satisfactory_subject_names = meta['satisfactory_subjects'].split(',')
            satisfactory_subjects = [s for s in subjects if s.name in satisfactory_subject_names]
        
        # Определяем, является ли класс "A" классом
        is_class_a = "А" in student.class_name or "A" in student.class_name
        
        # Находим заголовок "Уведомление" и добавляем основной текст после него
        header_found = False
        header_index = -1
        
        for i, paragraph in enumerate(doc.paragraphs):
            if "Уведомление" in paragraph.text:
                header_found = True
                header_index = i
                break
        
        if header_found:
            # Формируем основной текст уведомления
            main_text = f"Администрация школы доводит до Вашего сведения, что учащийся {student.class_name} класса {student.full_name} по результатам промежуточной аттестации имеет:"
            
            # Добавляем текст после заголовка
            if header_index + 1 < len(doc.paragraphs):
                # Если есть параграф после заголовка, используем его
                main_paragraph = doc.paragraphs[header_index + 1]
                main_paragraph.text = main_text
            else:
                # Если нет параграфа после заголовка, создаем новый
                main_paragraph = doc.add_paragraph(main_text)
            
            # Настраиваем форматирование: красная строка и выравнивание по ширине
            main_paragraph.paragraph_format.first_line_indent = Pt(1.25)  # Исправленный размер отступа
            main_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Продолжаем добавлять текст с академическими задолженностями
            insertion_index = header_index + 2
            
            # Добавляем список предметов с неудовлетворительными оценками
            if failed_subjects:
                failed_subjects_text = "; ".join([s.name for s in failed_subjects])  # Используем точку с запятой
                
                if insertion_index < len(doc.paragraphs):
                    doc.paragraphs[insertion_index].text = f"- неудовлетворительные отметки по предметам: {failed_subjects_text}"
                    insertion_index += 1
                else:
                    p = doc.add_paragraph(f"- неудовлетворительные отметки по предметам: {failed_subjects_text}")
                    insertion_index += 1
            
            # Для классов, кроме "А", также добавляем удовлетворительные оценки по углубленным предметам
            if not is_class_a and satisfactory_subjects:
                satisfactory_subjects_text = "; ".join([s.name for s in satisfactory_subjects])  # Используем точку с запятой
                
                if insertion_index < len(doc.paragraphs):
                    doc.paragraphs[insertion_index].text = f"- удовлетворительные отметки по предметам, изучаемым на углубленном уровне: {satisfactory_subjects_text}"
                    insertion_index += 1
                else:
                    p = doc.add_paragraph(f"- удовлетворительные отметки по предметам, изучаемым на углубленном уровне: {satisfactory_subjects_text}")
                    insertion_index += 1
            
            # Добавляем информацию о возможном переводе (для всех классов, кроме "А")
            if not is_class_a and 'deadline_date' in meta:
                deadline_date = meta['deadline_date']
                # Преобразуем формат даты в дд.мм.гггг
                try:
                    date_parts = deadline_date.split('-')
                    if len(date_parts) == 3:
                        formatted_date = f"{date_parts[2]}.{date_parts[1]}.{date_parts[0]}"
                    else:
                        formatted_date = deadline_date
                except:
                    formatted_date = deadline_date
                
                deadline_text = f"В случае, если данные оценки не будут исправлены до {formatted_date}, то по решению Педагогического совета в соответствии с положением школы о классах с углубленным изучением отдельных предметов может быть осуществлен перевод в общеобразовательный класс."
                
                # Создаем новый параграф с правильным форматированием
                p = doc.add_paragraph(deadline_text)
                p.paragraph_format.first_line_indent = Pt(1.25)  # Добавляем красную строку
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                insertion_index += 1
            
            # Добавляем график ликвидации задолженностей, если он есть
            if consultations:
                # Создаем параграф для заголовка графика
                p = doc.add_paragraph("График ликвидации академической задолженности:")
                insertion_index += 1
                
                # Создаем таблицу для графика
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                # Заголовки таблицы
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Предмет"
                hdr_cells[1].text = "Тема"
                hdr_cells[2].text = "Дата"
                hdr_cells[3].text = "Время"
                
                # Группируем консультации по предметам
                consultations_by_subject = {}
                for consultation in consultations:
                    subject_name = consultation['subject_name']
                    if subject_name not in consultations_by_subject:
                        consultations_by_subject[subject_name] = []
                    
                    consultations_by_subject[subject_name].append(consultation)
                
                # Добавляем строки для каждого предмета
                for subject_name, subject_consultations in consultations_by_subject.items():
                    for i, consultation in enumerate(subject_consultations):
                        row_cells = table.add_row().cells
                        
                        # Для первой консультации предмета указываем название предмета
                        if i == 0:
                            row_cells[0].text = subject_name
                        else:
                            row_cells[0].text = ""
                        
                        row_cells[1].text = consultation['topic_name'] or ""
                        
                        # Преобразуем формат даты в дд.мм.гггг
                        date_str = consultation['date']
                        try:
                            date_parts = date_str.split('-')
                            if len(date_parts) == 3:
                                formatted_date = f"{date_parts[2]}.{date_parts[1]}.{date_parts[0]}"
                            else:
                                formatted_date = date_str
                        except:
                            formatted_date = date_str
                        
                        row_cells[2].text = formatted_date
                        row_cells[3].text = consultation['time']
                
                # Добавляем пустую строку после таблицы
                doc.add_paragraph("")
            
        else:
            # Если заголовок не найден, добавляем его и весь текст с нуля
            # Но этого не должно происходить, если шаблон верный
            return {'success': False, 'message': 'В шаблоне не найден заголовок "Уведомление"'}
        
        # Создаем директорию для сохранения результатов, если её нет
        output_dir = 'generated_documents'
        os.makedirs(output_dir, exist_ok=True)
        
        # Формируем имя файла
        current_time = datetime.now().strftime('%Y%m%d_%H%M%S')
        file_name = f"{student.full_name}_{student.class_name}_{notification_data['period']}_{current_time}.docx"
        file_path = os.path.join(output_dir, file_name)
        
        # Сохраняем документ
        doc.save(file_path)
        
        return {
            'success': True, 
            'message': 'Документ успешно создан', 
            'file_path': file_path,
            'file_name': file_name
        }
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {'success': False, 'message': f'Ошибка при генерации документа: {str(e)}'}